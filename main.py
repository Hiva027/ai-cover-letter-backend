from dotenv import load_dotenv
load_dotenv()
import os
import tempfile
import asyncio
import logging
from fastapi import FastAPI, UploadFile, Form, Request
from fastapi.responses import JSONResponse, FileResponse
from fastapi.middleware.cors import CORSMiddleware
from openai import OpenAI
from pdf2docx import Converter
from docx import Document
from fpdf import FPDF
import tempfile, os, textwrap
from PyPDF2 import PdfReader

app = FastAPI()
origins = [
    "https://ai-cover-letter-flax.vercel.app",  # your frontend domain
    "http://localhost:3000",  # optional for local testing
]

app.add_middleware(
    CORSMiddleware,
    allow_origins=origins,
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)



client = OpenAI(api_key=os.getenv("OPENAI_API_KEY"))

# --- Text extraction helpers ---
def extract_text_from_docx(path):
    doc = Document(path)
    return "\n".join(p.text for p in doc.paragraphs if p.text.strip())

def extract_text_from_pdf(path):
    try:
        temp_docx = path.replace(".pdf", ".docx")
        cv = Converter(path)
        cv.convert(temp_docx)
        cv.close()
        text = extract_text_from_docx(temp_docx)
        os.remove(temp_docx)
        return text
    except Exception as e:
        print("[WARNING] pdf2docx failed, fallback to PyPDF2:", e)
        reader = PdfReader(path)
        text = "\n".join(page.extract_text() or "" for page in reader.pages)
        return text

# --- Cover letter generation ---
@app.post("/generate_cover_letter")
async def generate_cover_letter(
    resume: UploadFile,
    job_description: str = Form(""),
    word_count: int = Form(200)
):
    tmp_path = None
    try:
        with tempfile.NamedTemporaryFile(delete=False, suffix=f"_{resume.filename}") as tmp:
            tmp.write(await resume.read())
            tmp_path = tmp.name
        print(f"[INFO] 📄 Temp file created: {tmp_path}")

        resume_text = (
            extract_text_from_pdf(tmp_path)
            if resume.filename.lower().endswith(".pdf")
            else extract_text_from_docx(tmp_path)
        )

        prompt = f"""
        You are a professional career assistant. Write a plagiarism-free, formal cover letter 
        based on the following resume and job description.

        Guidelines:
        - Keep it under {word_count} words.
        - Maintain a standard 3–4 paragraph format.
        - Focus on achievements and tone relevant to the role.

        Resume:
        {resume_text}

        Job Description:
        {job_description}
        """

        completion = client.chat.completions.create(
            model="gpt-4o-mini",
            messages=[
                {"role": "system", "content": "You are a helpful assistant that writes professional cover letters."},
                {"role": "user", "content": prompt},
            ],
            timeout=60,
        )

        cover_letter = completion.choices[0].message.content.strip()
        print("[INFO] ✅ Cover letter generated successfully")
        return JSONResponse({"cover_letter": cover_letter})

    except Exception as e:
        print("[ERROR] ❌ Generation failed:", e)
        return JSONResponse({"error": str(e)}, status_code=500)
    finally:
        if tmp_path and os.path.exists(tmp_path):
            try:
                os.remove(tmp_path)
                print(f"[INFO] 🧹 Deleted temp file: {tmp_path}")
            except Exception as e:
                print(f"[WARNING] ⚠️ Could not delete temp file: {e}")

# --- Download DOCX ---
@app.post("/download_docx")
async def download_docx(request: Request):
    data = await request.json()
    text = data.get("text", "").strip()
    tmp = tempfile.NamedTemporaryFile(delete=False, suffix=".docx")
    doc = Document()
    for line in text.split("\n"):
        doc.add_paragraph(line)
    doc.save(tmp.name)
    return FileResponse(tmp.name, filename="cover_letter.docx")

# --- Download PDF (fixed) ---
@app.post("/download_pdf")
async def download_pdf(request: Request):
    try:
        data = await request.json()
        text = data.get("text", "").strip()
        if not text:
            return JSONResponse({"error": "No text provided"}, status_code=400)

        tmp = tempfile.NamedTemporaryFile(delete=False, suffix=".pdf")

        pdf = FPDF()
        pdf.add_page()
        pdf.set_font("Arial", size=12)
        pdf.set_auto_page_break(auto=True, margin=15)

        for paragraph in text.split("\n"):
            paragraph = paragraph.strip()
            if not paragraph:
                pdf.ln(5)
                continue
            wrapped = textwrap.fill(paragraph, width=90)
            safe_text = wrapped.encode("latin-1", "replace").decode("latin-1")
            pdf.multi_cell(0, 10, safe_text)

        pdf.output(tmp.name)
        return FileResponse(tmp.name, filename="cover_letter.pdf")

    except Exception as e:
        print("[ERROR] PDF generation failed:", e)
        return JSONResponse({"error": str(e)}, status_code=500)


# ------------------- Test Endpoint -------------------
@app.get("/test_openai")
async def test_openai():
    try:
        response = client.chat.completions.create(
            model="gpt-4o-mini",
            messages=[
                {"role": "system", "content": "You are a helpful assistant."},
                {"role": "user", "content": "Write a short greeting message to confirm API is working."},
            ],
        )
        return {"message": response.choices[0].message.content}
    except Exception as e:
        logging.exception("❌ Error testing OpenAI API")
        return {"error": str(e)}
